"""
Multi-format report generation for vulnscan results.

Supports: Markdown, JSON, HTML, DOCX (VAPT-grade).
DOCX output is only available when python-docx is installed.
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from typing import Any, Optional


# ── Internal helpers ──────────────────────────────────────────────────────────

_RISK_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "MINIMAL": 4, "NONE": 5}


def _risk_sort_key(result: dict) -> tuple[int, int]:
    level = (result.get("risk_level") or "NONE").upper()
    return _RISK_ORDER.get(level, 99), -(result.get("risk_score") or 0)


def _severity_badge(level: str) -> str:
    badges = {
        "CRITICAL": "🔴 CRITICAL",
        "HIGH": "🟠 HIGH",
        "MEDIUM": "🟡 MEDIUM",
        "LOW": "🟢 LOW",
        "MINIMAL": "⬜ MINIMAL",
        "NONE": "— NONE",
    }
    return badges.get(level.upper(), level)


# ── VulnReporter ─────────────────────────────────────────────────────────────

class VulnReporter:
    """Generates scan reports in multiple formats from a unified results dict."""

    def __init__(self, results: dict[str, Any]):
        """
        Args:
            results: Dict with keys:
                scan_date, total_targets, total_services, total_cves,
                scan_metadata (optional), results (list of service dicts)
        """
        self.results = results
        self.sorted_findings: list[dict] = sorted(
            results.get("results", []),
            key=_risk_sort_key,
        )

    # ── JSON ──────────────────────────────────────────────────────────────────

    def export_json(self, path: str) -> str:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w") as f:
            json.dump(self.results, f, indent=2, default=str)
        return path

    # ── Markdown ─────────────────────────────────────────────────────────────

    def export_markdown(self, path: str) -> str:
        lines = [
            "# Network Vulnerability Report\n",
            f"**Scan Date:** {self.results.get('scan_date', 'N/A')}  ",
            f"**Targets:** {self.results.get('total_targets', 0)} | "
            f"**Services:** {self.results.get('total_services', 0)} | "
            f"**CVEs:** {self.results.get('total_cves', 0)}\n",
        ]

        meta = self.results.get("scan_metadata", {})
        if meta:
            lines += [
                "## Scan Metadata\n",
                f"- **Operator:** {meta.get('operator', 'N/A')}",
                f"- **Ticket:** {meta.get('ticket', 'N/A')}",
                f"- **Engagement:** {meta.get('engagement', 'N/A')}",
                f"- **Profile:** {meta.get('profile', 'N/A')}\n",
            ]

        lines.append("## Findings\n")
        for r in self.sorted_findings:
            ip = r.get("target_ip", "")
            port = r.get("port", "")
            product = f"{r.get('product', '')} {r.get('version', '')}".strip() or r.get("service_name", "unknown")
            risk = _severity_badge(r.get("risk_level", "NONE"))
            score = r.get("risk_score", 0)

            lines += [
                f"### {ip}:{port} — {product}",
                f"**Risk:** {risk} ({score}/100)\n",
            ]

            cves = r.get("cve_details", [])
            if cves:
                lines.append("| CVE | CVSS | Severity | Description |")
                lines.append("|-----|------|----------|-------------|")
                for cve in cves:
                    desc = (cve.get("description") or "")[:100].replace("|", "\\|")
                    lines.append(
                        f"| {cve.get('id')} | {cve.get('cvss_score')} "
                        f"| {cve.get('severity')} | {desc}... |"
                    )
                lines.append("")

            ai = (r.get("ai_analysis") or "").strip()
            if ai and ai != "AI analysis disabled":
                lines += ["**AI Analysis:**\n", ai, ""]

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w") as f:
            f.write("\n".join(lines))
        return path

    # ── HTML ─────────────────────────────────────────────────────────────────

    def export_html(self, path: str) -> str:
        RISK_COLORS = {
            "CRITICAL": "#dc2626", "HIGH": "#ea580c", "MEDIUM": "#d97706",
            "LOW": "#16a34a", "MINIMAL": "#6b7280", "NONE": "#9ca3af",
        }

        rows_html = ""
        for r in self.sorted_findings:
            level = (r.get("risk_level") or "NONE").upper()
            color = RISK_COLORS.get(level, "#6b7280")
            product = f"{r.get('product', '')} {r.get('version', '')}".strip() or r.get("service_name", "—")
            cve_count = len(r.get("cve_details") or [])
            rows_html += (
                f"<tr>"
                f"<td>{r.get('target_ip', '')}</td>"
                f"<td>{r.get('port', '')}</td>"
                f"<td>{r.get('service_name', '')}</td>"
                f"<td>{product}</td>"
                f"<td style='color:{color};font-weight:bold'>{level} ({r.get('risk_score', 0)}/100)</td>"
                f"<td>{cve_count}</td>"
                f"</tr>\n"
            )

        cve_detail_html = ""
        for r in self.sorted_findings:
            cves = r.get("cve_details") or []
            if not cves:
                continue
            ip = r.get("target_ip", "")
            port = r.get("port", "")
            product = f"{r.get('product', '')} {r.get('version', '')}".strip()
            cve_detail_html += f"<h3>{ip}:{port} — {product}</h3><table border='1'>"
            cve_detail_html += "<tr><th>CVE</th><th>CVSS</th><th>Severity</th><th>Description</th></tr>"
            for cve in cves:
                cve_detail_html += (
                    f"<tr><td>{cve.get('id')}</td><td>{cve.get('cvss_score')}</td>"
                    f"<td>{cve.get('severity')}</td><td>{cve.get('description', '')}</td></tr>"
                )
            cve_detail_html += "</table><br>"

        html = f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>Vulnerability Report</title>
<style>
  body{{font-family:Arial,sans-serif;margin:40px;background:#f9fafb}}
  h1,h2{{color:#1f2937}} table{{border-collapse:collapse;width:100%;margin-bottom:20px}}
  th{{background:#1f2937;color:white;padding:8px}} td{{padding:8px;border:1px solid #d1d5db}}
  tr:nth-child(even){{background:#f3f4f6}}
  .meta{{background:#e5e7eb;padding:16px;border-radius:8px;margin-bottom:20px}}
</style></head>
<body>
<h1>Network Vulnerability Report</h1>
<div class="meta">
  <b>Scan Date:</b> {self.results.get('scan_date', 'N/A')} &nbsp;|&nbsp;
  <b>Targets:</b> {self.results.get('total_targets', 0)} &nbsp;|&nbsp;
  <b>Services:</b> {self.results.get('total_services', 0)} &nbsp;|&nbsp;
  <b>CVEs:</b> {self.results.get('total_cves', 0)}
</div>
<h2>Findings Summary</h2>
<table>
<tr><th>Host</th><th>Port</th><th>Service</th><th>Product</th><th>Risk</th><th>CVEs</th></tr>
{rows_html}
</table>
<h2>CVE Details</h2>
{cve_detail_html or "<p>No CVEs found.</p>"}
</body></html>"""

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w") as f:
            f.write(html)
        return path

    # ── DOCX (VAPT-grade) ────────────────────────────────────────────────────

    def export_docx(
        self,
        path: str,
        org_name: str = "Organization",
        group_name: str = "Security Team",
        unit_name: str = "Internal VAPT",
        report_title: str = "Network Vulnerability Assessment & Penetration Testing",
        classification: str = "INTERNAL",
    ) -> str:
        """
        Generate an evidence-grade DOCX report.
        Requires: pip install python-docx
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Cover page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(report_title)
        run.bold = True
        run.font.size = Pt(20)

        doc.add_paragraph()

        meta_table = doc.add_table(rows=0, cols=2)
        meta_table.style = "Table Grid"
        for k, v in {
            "Organization": org_name,
            "Group": group_name,
            "Unit": unit_name,
            "Classification": classification,
            "Generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Scan Date": str(self.results.get("scan_date", "")),
            "Total Targets": str(self.results.get("total_targets", 0)),
            "Total Services": str(self.results.get("total_services", 0)),
            "Total CVEs": str(self.results.get("total_cves", 0)),
        }.items():
            row = meta_table.add_row().cells
            row[0].text = k
            row[1].text = v

        # Scan metadata (VAPT audit trail)
        meta = self.results.get("scan_metadata", {})
        if meta:
            doc.add_paragraph()
            audit_table = doc.add_table(rows=0, cols=2)
            audit_table.style = "Table Grid"
            for k, v in {
                "Operator": meta.get("operator", "N/A"),
                "Ticket / Change Ref": meta.get("ticket", "N/A"),
                "Engagement ID": meta.get("engagement", "N/A"),
                "Scan Profile": meta.get("profile", "N/A"),
                "Notes": meta.get("notes", ""),
            }.items():
                row = audit_table.add_row().cells
                row[0].text = k
                row[1].text = str(v)

        doc.add_page_break()

        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(
            "This report summarises the results of an internal network vulnerability assessment "
            "conducted to identify exposed services, known CVEs, and risk exposure. "
            "Findings are prioritised by risk score for remediation planning."
        )

        doc.add_heading("2. Scope & Methodology", level=1)
        for bullet in [
            "Network port/service discovery using Nmap with service version detection.",
            "CVE enrichment via NVD API (CPE-based primary, keyword fallback).",
            "Risk scoring using absolute CVSS weight accumulation (not ratio-normalised).",
            "Exploit reference lookup via cve.circl.lu and GitHub.",
            "AI-assisted remediation guidance (advisory only — requires human validation).",
        ]:
            doc.add_paragraph(f"• {bullet}")

        doc.add_heading("3. Findings Overview", level=1)
        summary = doc.add_table(rows=1, cols=7)
        summary.style = "Table Grid"
        hdr = summary.rows[0].cells
        for i, h in enumerate(["Host", "Port", "Service", "Product", "Version", "Risk", "CVEs"]):
            hdr[i].text = h

        for r in self.sorted_findings[:50]:
            row = summary.add_row().cells
            row[0].text = str(r.get("target_ip", ""))
            row[1].text = str(r.get("port", ""))
            row[2].text = str(r.get("service_name", ""))
            row[3].text = str(r.get("product", ""))
            row[4].text = str(r.get("version", ""))
            row[5].text = f"{r.get('risk_level', 'NONE')} ({r.get('risk_score', 0)}/100)"
            row[6].text = str(len(r.get("cve_details") or []))

        doc.add_page_break()
        doc.add_heading("4. Detailed Findings", level=1)

        for r in self.sorted_findings:
            ip = r.get("target_ip", "")
            port = r.get("port", "")
            product = f"{r.get('product', '')} {r.get('version', '')}".strip()
            doc.add_heading(f"{ip}:{port} — {product or r.get('service_name', 'unknown')}", level=2)

            info = doc.add_table(rows=0, cols=2)
            info.style = "Table Grid"
            for k, v in {
                "Target IP": ip, "Port": str(port),
                "Service": r.get("service_name", ""),
                "Product": r.get("product", ""),
                "Version": r.get("version", ""),
                "Risk Level": f"{r.get('risk_level', 'NONE')} ({r.get('risk_score', 0)}/100)",
            }.items():
                row = info.add_row().cells
                row[0].text = k
                row[1].text = str(v)

            cves = r.get("cve_details") or []
            doc.add_paragraph("CVE Evidence:")
            if cves:
                cve_tbl = doc.add_table(rows=1, cols=4)
                cve_tbl.style = "Table Grid"
                h = cve_tbl.rows[0].cells
                h[0].text = "CVE ID"
                h[1].text = "CVSS"
                h[2].text = "Severity"
                h[3].text = "Description"
                for cve in cves:
                    row = cve_tbl.add_row().cells
                    row[0].text = str(cve.get("id", ""))
                    row[1].text = str(cve.get("cvss_score", "N/A"))
                    row[2].text = str(cve.get("severity", "UNKNOWN"))
                    row[3].text = str(cve.get("description", ""))
            else:
                doc.add_paragraph("No CVEs identified from current enrichment sources.")

            ai_text = (r.get("ai_analysis") or "").strip()
            doc.add_paragraph("AI-Assisted Analysis (Advisory):")
            doc.add_paragraph(
                ai_text if ai_text and ai_text != "AI analysis disabled"
                else "AI analysis not run or unavailable."
            )
            doc.add_paragraph("")

        doc.add_heading("Appendix A — Notes", level=1)
        doc.add_paragraph(
            "AI-assisted analysis is advisory and must be validated by engineering owners "
            "before implementation. CVE data sourced from NVD (nvd.nist.gov)."
        )

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)
        return path
